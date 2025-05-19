import tkinter as tk
from tkinter import ttk, filedialog, messagebox
from docx import Document
from docx.opc.exceptions import PackageNotFoundError
import os
from pathlib import Path
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from docx.oxml.ns import qn
from docxtpl import DocxTemplate


class InquiryGenerator:
    def __init__(self):
        self.root = tk.Tk()
        self.root.title("海关询问笔录生成系统 v3.3")
        self.root.geometry("1100x1000")

        # 初始化界面
        self.setup_ui()

        # 初始化模板路径
        self.template_path = None
        self.locate_template()

    def setup_ui(self):
        """创建用户界面"""
        main_frame = ttk.Frame(self.root, padding=20)
        main_frame.pack(fill=tk.BOTH, expand=True)

        # 状态栏
        self.status_var = tk.StringVar()
        ttk.Label(main_frame, textvariable=self.status_var, foreground="blue").pack(side=tk.TOP, fill=tk.X)

        # 被询问人信息区
        info_frame = ttk.LabelFrame(main_frame, text="被询问人信息", padding=10)
        info_frame.pack(fill=tk.X, pady=10)

        # button1 = ttk.Button(info_frame, text="按钮1")
        # button1.grid(row=0, column=0, padx=5, pady=5)

        fields = [
            ("被询问人", 0, 0), ("性别", 0, 1),
            ("证件名称", 0, 2), ("证件号码", 1, 0),
            ("住址", 1, 1), ("年龄", 1, 2),
            ("工作单位", 3, 0), ("单位地址", 3, 1),
            ("询问地点", 4, 1),
            ("职务", 3, 2), ("联系方式", 4, 0),
            ("询问时间起", 5, 0), ("询问时间止", 5, 1),
            ("询问单位", 4, 2)
        ]

        # keys_location = {
        #     "被询问人": (0, 1), "性别": (0, 3), "证件名称": (1, 1), "证件号码": (1, 3),
        #     "住址": (2, 1), "年龄": (2, 3), "工作单位": (3, 1), "单位地址": (3, 3),
        #     "职务": (4, 1), "联系方式": (4, 3), "询问时间起": (5, 1), "询问时间止": (5, 3),
        #     "询问地点": (6, 1)
        # }

        self.entries = {}
        for label, row, col in fields:
            frame = ttk.Frame(info_frame)
            frame.grid(row=row, column=col, sticky="ew", padx=5, pady=3)

            ttk.Label(frame, text=label + ":", width=9).pack(side=tk.LEFT)

            if label == "性别":
                entry = ttk.Combobox(frame, values=["男", "女"], width=15)
            elif "时间" in label:
                entry = ttk.Entry(frame, width=40, )
                ttk.Label(frame).pack(side=tk.LEFT)
            elif "询问单位" in label:
                entry = ttk.Entry(frame, width=25)
                entry.insert(0, "中华人民共和国厦门海关")
            else:
                entry = ttk.Entry(frame, width=25)

            entry.pack(side=tk.LEFT)
            self.entries[label] = entry

        # 询问内容区（带滚动条）
        qa_frame = ttk.LabelFrame(main_frame, text="询问内容", padding=10)
        qa_frame.pack(fill=tk.BOTH, expand=True, pady=10)

        self.qa_canvas = tk.Canvas(qa_frame, borderwidth=0, highlightthickness=0)
        vsb = ttk.Scrollbar(qa_frame, orient="vertical", command=self.qa_canvas.yview)
        self.qa_container = ttk.Frame(self.qa_canvas)

        self.qa_canvas.configure(yscrollcommand=vsb.set)
        self.qa_container.bind("<Configure>",
                               lambda e: self.qa_canvas.configure(scrollregion=self.qa_canvas.bbox("all"))
                               )
        vsb.pack(side="right", fill="y")
        self.qa_canvas.pack(side="left", fill="both", expand=True)
        self.qa_canvas.create_window((0, 0), window=self.qa_container, anchor="nw")

        self.qa_rows = []
        self.add_qa_row()

        # 控制按钮
        btn_frame = ttk.Frame(main_frame)
        btn_frame.pack(fill=tk.X, pady=10)
        ttk.Button(btn_frame, text="+ 添加问答", command=self.add_qa_row).pack(side=tk.LEFT, padx=5)
        ttk.Button(btn_frame, text="生成笔录", command=self.generate_document).pack(side=tk.RIGHT)

    def add_qa_row(self):
        """添加问答对"""
        frame = ttk.Frame(self.qa_container)
        frame.pack(fill=tk.X, pady=5)

        # 问
        q_frame = ttk.Frame(frame)
        q_frame.pack(fill=tk.X, pady=1)
        ttk.Label(q_frame, text=f"问:").pack(side=tk.LEFT)
        q_entry = tk.Text(q_frame, height=3, width=125, wrap=tk.WORD)
        q_entry.pack(side=tk.LEFT, padx=5, fill=tk.X, expand=True)

        # 答
        a_frame = ttk.Frame(frame)
        a_frame.pack(fill=tk.X, pady=2)
        ttk.Label(a_frame, text="答:").pack(side=tk.LEFT)
        a_entry = tk.Text(a_frame, height=3, width=125, wrap=tk.WORD)
        a_entry.pack(side=tk.LEFT, padx=5, fill=tk.X, expand=True)

        self.qa_rows.append((q_entry, a_entry))
        self.qa_canvas.yview_moveto(1)  # 滚动到底部

    def locate_template(self):
        """查找并加载模板文件"""
        # 自动查找
        template_path = self.auto_find_template()

        if not template_path:
            # 手动选择模板
            template_path = filedialog.askopenfilename(
                filetypes=[("Word模板", "*.docx")],
                title="选择询问笔录模板文件"
            )

        if template_path:
            self.template_path = template_path
            self.update_status(f"已加载模板: {os.path.basename(template_path)}")
        else:
            messagebox.showwarning("警告", "未选择模板文件，请通过菜单重新加载")

    def auto_find_template(self):
        """自动查找常见位置的模板文件"""
        search_paths = [
            Path("询问笔录模板.doc"),
            Path("询问笔录模板.docx"),
            Path(__file__).parent / "templates/询问笔录模板.docx",
            Path.home() / "Desktop/询问笔录模板.docx"
        ]
        for path in search_paths:
            if path.exists():
                return str(path)
        return None

    def update_status(self, message):
        """更新状态栏"""
        self.status_var.set(message)

    def generate_document(self):
        """生成询问笔录文档"""
        if not self.template_path:
            messagebox.showerror("错误", "未加载模板文件")
            return

        try:
            # 验证必填字段
            required_fields = ["被询问人", "性别", "证件号码", "询问时间起", "询问时间止", "询问地点"]
            for field in required_fields:
                if not self.entries[field].get().strip():
                    messagebox.showwarning("输入不完整", f"请填写【{field}】字段")
                    return

            # 加载模板
            doc = Document(self.template_path)

            # 填充基本信息到表格
            for table in doc.tables:
                for row_index, row in enumerate(table.rows):
                    for cell_index, cell in enumerate(row.cells):
                        text = cell.text
                        if text in self.entries:
                            row.cells[cell_index + 1].text = f"{self.entries[text].get()}"

            # 单独修改时间
            doc.tables[0].rows[5].cells[
                1].text = f"{self.entries["询问时间起"].get()}至{self.entries["询问时间止"].get()}"

            # 单独修改表头
            title_context = {"title": self.entries["询问单位"].get()}

            doc_tpl = DocxTemplate(self.template_path)
            doc_tpl.render(title_context)

            # 插入问答内容
            # for q, a in self.qa_rows:

            # 插入问答内容
            # 获取目标表格
            table = doc.tables[0]  # 选择你需要操作的表格

            # 添加新行并合并为一个单元格
            new_row = table.add_row()
            num_cells = len(new_row.cells)
            merged_cell = new_row.cells[0]
            for i in range(1, num_cells):
                merged_cell = merged_cell.merge(new_row.cells[i])

            # 清空合并单元格中默认的段落内容
            merged_cell._element.clear_content()

            # 在合并后的单元格中逐段添加问答
            for q, a in self.qa_rows:
                q_text = q.get("1.0", tk.END).strip()
                a_text = a.get("1.0", tk.END).strip()

                if q_text:
                    para_q = merged_cell.add_paragraph()
                    para_q.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                    run_q = para_q.add_run(f"问：{q_text}")
                    run_q.font.name = '宋体'
                    run_q.font.size = Pt(16)
                    run_q.underline = True
                    run_q._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

                if a_text:
                    para_a = merged_cell.add_paragraph()
                    para_a.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                    run_a = para_a.add_run(f"答：{a_text}")
                    run_a.font.name = '宋体'
                    run_a.font.size = Pt(16)
                    run_a.underline = True
                    run_a._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

            # start, end = 0, len(self.qa_rows)
            # for row_index in range(7, len(doc.tables[0].rows), 2):
            #
            #     if start < end:
            #         doc.tables[0].rows[row_index].cells[0].text =  self.qa_rows[start][0].get("1.0", tk.END).strip()
            #         doc.tables[0].rows[row_index + 1].cells[0].text = self.qa_rows[start][1].get("1.0", tk.END).strip()
            #         start += 1
            #         continue

            #
            # q_text = q.get("1.0", tk.END).strip()
            # a_text = a.get("1.0", tk.END).strip()
            #
            # if q_text:
            #     new_q = doc.add_paragraph(f"问：{q_text}")
            #     new_q.runs[0].bold = True
            # if a_text:
            #     doc.add_paragraph(f"答：{a_text}")
            # doc.add_paragraph()

            # 保存文件
            output_path = filedialog.asksaveasfilename(
                defaultextension=".docx",
                filetypes=[("Word文档", "*.docx")],
                initialfile="海关询问笔录_已填写.docx"
            )
            if output_path:
                doc.save(output_path)
                messagebox.showinfo("成功", f"文档已保存到:\n{output_path}")
                os.startfile(output_path)  # 自动打开文件

                # 单独修改表头
                title_context = {"title": self.entries["询问单位"].get()}
                doc_tpl = DocxTemplate(output_path)
                doc_tpl.render(title_context)
                doc_tpl.save(output_path)





        except PackageNotFoundError:
            messagebox.showerror("错误", "模板文件损坏或不是有效的.docx文件")
        except Exception as e:
            messagebox.showerror("错误", f"生成文档时出错:\n{str(e)}")

    def run(self):
        self.root.mainloop()


if __name__ == "__main__":
    app = InquiryGenerator()
    app.run()
